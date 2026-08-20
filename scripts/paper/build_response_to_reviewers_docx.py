#!/usr/bin/env python3
"""Build a submission-ready Word response from the canonical Markdown file."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = "https://github.com/SoroushVahidi/Augmented-caching"
MD_REL_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
MATH_RE = re.compile(r"\$([^$]+)\$")


def latex_to_unicode(s: str) -> str:
    repl = {
        r"\pi_1": "π₁",
        r"\pi_2": "π₂",
        r"\mu": "μ",
        r"\times": "×",
        r"\approx": "≈",
        r"\to": "→",
        r"\,": " ",
        r"\ ": " ",
        r"{,}": ",",
        r"\%": "%",
        r"\emph{": "",
    }
    out = s
    for a, b in repl.items():
        out = out.replace(a, b)
    out = out.replace("{", "").replace("}", "")
    return out


def clean_inline(text: str) -> str:
    text = MATH_RE.sub(lambda m: latex_to_unicode(m.group(1)), text)
    text = text.replace("—", "—").replace("–", "–")
    return text


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, text: str, *, italic: bool = False, bold: bool = False) -> None:
    text = clean_inline(text)
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://[^\s)]+)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.italic = italic
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = italic
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.italic = italic
            run.bold = bold
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("["):
            label, url = MD_REL_RE.match(token).groups()
            if url.startswith("http"):
                abs_url = url
            else:
                rel = url.lstrip("./")
                abs_url = f"{REPO}/blob/main/{rel}"
            add_hyperlink(paragraph, label, abs_url)
        else:
            add_hyperlink(paragraph, token.rstrip(".,;"), token.rstrip(".,;"))
            trail = token[len(token.rstrip(".,;")) :]
            if trail:
                run = paragraph.add_run(trail)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.italic = italic
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element.get_or_add_rPr()
    rfonts = r.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        r.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def shade_cell(cell, fill: str) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcpr.append(shd)


def set_cell_margins(cell, **twips) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = OxmlElement("w:tcMar")
    for edge, val in twips.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcmar.append(node)
    tcpr.append(tcmar)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    set_run_font(run, size=9)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)
    set_run_font(r2, size=9)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if set(line.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_runs(p, text, bold=(i == 0))
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            if i == 0:
                shade_cell(cell, "F2F2F2")
                for run in p.runs:
                    run.bold = True


def style_quote_paragraph(p) -> None:
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "666666")
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F7F7F7")
    pPr.append(shd)


def build(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        st.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)

    i = 0
    quote_buf: list[str] = []
    table_buf: list[str] = []

    def flush_quote():
        nonlocal quote_buf
        if not quote_buf:
            return
        p = doc.add_paragraph()
        style_quote_paragraph(p)
        add_runs(p, " ".join(quote_buf), italic=True)
        quote_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        add_table(doc, parse_table(table_buf))
        doc.add_paragraph()
        table_buf = []

    while i < len(lines):
        raw = lines[i]
        if raw.startswith("|"):
            flush_quote()
            table_buf.append(raw)
            i += 1
            continue
        else:
            flush_table()

        if raw.startswith("> "):
            quote_buf.append(raw[2:])
            i += 1
            continue
        else:
            flush_quote()

        if raw.strip() == "---":
            i += 1
            continue
        if raw.startswith("# ") and not raw.startswith("##"):
            title = raw[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            set_run_font(run, size=16, bold=True)
            i += 1
            continue
        if raw.startswith("## "):
            h = doc.add_heading(raw[3:].strip(), level=1)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
        if raw.startswith("### "):
            h = doc.add_heading(raw[4:].strip(), level=2)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
        if raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, raw[2:])
            i += 1
            continue
        if not raw.strip():
            i += 1
            continue
        p = doc.add_paragraph()
        add_runs(p, raw)
        i += 1

    flush_quote()
    flush_table()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    root = Path(__file__).resolve().parents[2]
    md = root / "submission_kbs_revision_final" / "02_Response_to_Reviewers.md"
    out = root / "submission_kbs_revision_final" / "02_Response_to_Reviewers.docx"
    build(md, out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
